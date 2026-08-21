#!/usr/bin/env python3
"""Конвертер md-шаблонів кейсів → .docx для імпорту в Google Доки.
HTML-коментарі (<!-- -->) рендеряться як видимі службові примітки.
"""
import re, sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_runs(p, text, italic=False, color=None):
    # розбір **жирного**
    for i, chunk in enumerate(text.split("**")):
        if chunk == "":
            continue
        r = p.add_run(chunk)
        r.bold = (i % 2 == 1)
        r.italic = italic
        if color:
            r.font.color.rgb = color

GREY = RGBColor(0x80, 0x80, 0x80)
BLUE = RGBColor(0x29, 0x66, 0xFF)

def convert(md_path, docx_path):
    with open(md_path, encoding="utf-8") as f:
        lines = f.read().split("\n")
    doc = Document()
    # базовий шрифт
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(11)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]

        # --- HTML-коментар → службова примітка ---
        if line.strip().startswith("<!--"):
            block = []
            # зібрати до -->
            if "-->" in line:
                inner = line.strip()[4:].split("-->")[0]
                block = [inner]
                i += 1
            else:
                i += 1
                while i < n and "-->" not in lines[i]:
                    block.append(lines[i])
                    i += 1
                i += 1  # рядок з -->
            # відрендерити
            hdr = doc.add_paragraph()
            add_runs(hdr, "— Службова інструкція (не публікувати) —", italic=True, color=BLUE)
            for b in block:
                if b.strip() == "":
                    continue
                p = doc.add_paragraph()
                add_runs(p, b.rstrip(), italic=True, color=GREY)
            continue

        # --- таблиця ---
        if line.strip().startswith("|") and i + 1 < n and re.match(r"^\s*\|[\s:|-]+\|\s*$", lines[i+1] or ""):
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                rows.append(lines[i])
                i += 1
            header = [c.strip() for c in rows[0].strip().strip("|").split("|")]
            body = rows[2:]  # пропустити роздільник
            t = doc.add_table(rows=1, cols=len(header))
            t.style = "Light Grid Accent 1"
            for j, h in enumerate(header):
                cell = t.rows[0].cells[j]
                cell.paragraphs[0].text = ""
                add_runs(cell.paragraphs[0], h)
            for br in body:
                cells = [c.strip() for c in br.strip().strip("|").split("|")]
                row = t.add_row().cells
                for j in range(min(len(cells), len(header))):
                    row[j].paragraphs[0].text = ""
                    add_runs(row[j].paragraphs[0], cells[j])
            # опитувальник: 3 колонки з широкою «Відповідь» праворуч
            if len(header) == 3 and "Відповід" in header[2]:
                widths = [Inches(2.5), Inches(1.2), Inches(2.8)]
                t.autofit = False
                t.allow_autofit = False
                for r in t.rows:
                    for j, c in enumerate(r.cells):
                        c.width = widths[j]
            doc.add_paragraph()
            continue

        # --- заголовки ---
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3); i += 1; continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2); i += 1; continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1); i += 1; continue

        # --- горизонтальна лінія ---
        if line.strip() == "---":
            doc.add_paragraph("_" * 40); i += 1; continue

        # --- цитата/callout ---
        if line.strip().startswith(">"):
            p = doc.add_paragraph(style="Intense Quote")
            add_runs(p, line.strip().lstrip(">").strip())
            i += 1; continue

        # --- нумерований список ---
        m = re.match(r"^(\d+)\.\s+(.*)$", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_runs(p, m.group(2)); i += 1; continue

        # --- маркований список ---
        if re.match(r"^\s*-\s+", line):
            indent = len(line) - len(line.lstrip())
            p = doc.add_paragraph(style="List Bullet")
            if indent >= 2:
                p.paragraph_format.left_indent = Inches(0.5)
            add_runs(p, re.sub(r"^\s*-\s+", "", line)); i += 1; continue

        # --- порожній рядок ---
        if line.strip() == "":
            i += 1; continue

        # --- звичайний абзац ---
        p = doc.add_paragraph()
        add_runs(p, line)
        i += 1

    doc.save(docx_path)
    print("OK:", docx_path)

if __name__ == "__main__":
    pairs = [
        ("seo-ecommerce-template.md", "SEO-кейс-шаблон-магазин.docx"),
        ("seo-services-template.md",  "SEO-кейс-шаблон-послуги.docx"),
        ("intake-ecommerce.md",       "Опитувальник-кейс-магазин.docx"),
        ("intake-services.md",        "Опитувальник-кейс-послуги.docx"),
    ]
    for md, dx in pairs:
        convert(md, dx)
